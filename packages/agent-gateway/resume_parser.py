"""
简历解析器 — 文本提取 + LLM 结构化解析。

流程：
  1. 在 executor 中提取纯文本（PDF / DOCX）
  2. 调用 LLM 解析为结构化 JSON
  3. 更新 DB（status: parsing → done / failed）
"""
from __future__ import annotations

import asyncio
import json
import logging
import re
from concurrent.futures import ThreadPoolExecutor

from anthropic import AsyncAnthropic

import config
import resume_db

log = logging.getLogger(__name__)

_executor = ThreadPoolExecutor(max_workers=4, thread_name_prefix="resume-parse")

# ── 文本提取 ──────────────────────────────────────────────────────────────────

def _extract_pdf(file_path: str) -> str:
    from pdfminer.high_level import extract_text  # type: ignore
    return extract_text(file_path)


def _extract_docx(file_path: str) -> str:
    from docx import Document  # type: ignore
    doc = Document(file_path)
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("  ".join(cells))
    return "\n".join(parts)


def _clean_text(raw: str) -> str:
    # pdfminer 在 CJK 字体边界吐 \x00,postgres TEXT 不收 NUL byte。
    raw = raw.replace("\x00", "")
    lines = raw.splitlines()
    cleaned: list[str] = []
    for line in lines:
        line = line.strip()
        # 去掉纯数字页码行
        if re.fullmatch(r"\d+", line):
            continue
        if line:
            cleaned.append(line)
    text = "\n".join(cleaned)
    # 截断超长简历，保留前 8000 字
    return text[:8000]


async def _extract_text_async(file_path: str, file_type: str) -> str:
    loop = asyncio.get_event_loop()
    if file_type == "pdf":
        raw = await loop.run_in_executor(_executor, _extract_pdf, file_path)
    else:
        raw = await loop.run_in_executor(_executor, _extract_docx, file_path)
    return _clean_text(raw)


# ── LLM 解析 ──────────────────────────────────────────────────────────────────

_SYSTEM = (
    "你是简历解析专家。只返回 JSON，不要任何解释或 markdown 代码块。\n"
    "target_salary 统一转换为月薪人民币元整数（'15k'→15000，'面议'→null）。\n"
    "未找到的字段返回 null 或 []。"
)

_SCHEMA_EXAMPLE = """{
  "name": "张三",
  "phone": "138xxxx1234",
  "email": "zhangsan@example.com",
  "work_years": "5年",
  "degree": "本科",
  "target_salary_raw": "15k-20k",
  "target_salary_min": 15000,
  "target_salary_max": 20000,
  "target_cities": ["北京", "上海"],
  "target_positions": ["后端工程师", "Python开发"],
  "work_experience": [
    {"company": "公司名", "title": "职位", "duration": "2020.01-2023.06", "desc": "工作描述"}
  ],
  "education": [
    {"school": "学校名", "degree": "本科", "major": "计算机科学", "duration": "2015-2019"}
  ],
  "skills": ["Python", "Go", "PostgreSQL"],
  "projects": [
    {"name": "项目名", "duration": "2022.01-2022.06", "desc": "项目描述"}
  ],
  "self_evaluation": "自我评价文本"
}"""


def _normalize_salary(parsed: dict) -> dict:
    """将字符串薪资转换为整数（容错处理）。"""
    for key in ("target_salary_min", "target_salary_max"):
        val = parsed.get(key)
        if isinstance(val, str):
            m = re.search(r"(\d+)", val.replace(",", ""))
            if m:
                num = int(m.group(1))
                # 如果是 "k" 单位
                if "k" in val.lower() or num < 1000:
                    num *= 1000
                parsed[key] = num
            else:
                parsed[key] = None
    return parsed


def _extract_json_object(text: str) -> str:
    """从 LLM 响应中提取最外层 JSON 对象（括号匹配法）。"""
    depth = 0
    start = -1
    for i, ch in enumerate(text):
        if ch == '{':
            if depth == 0:
                start = i
            depth += 1
        elif ch == '}':
            depth -= 1
            if depth == 0 and start != -1:
                return text[start:i + 1]
    return text


def _fix_control_chars(text: str) -> str:
    """将 JSON 字符串值内的裸控制字符转义（LLM 常见输出问题）。"""
    result: list[str] = []
    in_str = False
    backslash_count = 0
    for ch in text:
        if ch == '\\':
            backslash_count += 1
            result.append(ch)
            continue
        escaped = (backslash_count % 2 == 1)
        backslash_count = 0
        if ch == '"' and not escaped:
            in_str = not in_str
            result.append(ch)
        elif in_str and ch == '\n':
            result.append('\\n')
        elif in_str and ch == '\r':
            result.append('\\r')
        elif in_str and ch == '\t':
            result.append('\\t')
        else:
            result.append(ch)
    return ''.join(result)


def _parse_llm_response(raw: str) -> dict:
    """多阶段 JSON 解析：直接解析 → 提取对象 → json_repair → 修复控制字符。"""
    # Stage 1: 去除 markdown 代码块
    text = raw.strip()
    text = re.sub(r"^```(?:json)?\s*\n?", "", text, flags=re.MULTILINE)
    text = re.sub(r"\n?```\s*$", "", text, flags=re.MULTILINE)
    text = text.strip()

    # Stage 2: 直接解析
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # Stage 3: 括号匹配提取最外层 JSON 对象后再解析
    obj = _extract_json_object(text)
    try:
        return json.loads(obj)
    except json.JSONDecodeError:
        pass

    # Stage 4: json_repair（处理未转义引号、截断、多余逗号等 LLM 常见问题）
    try:
        from json_repair import repair_json  # type: ignore
        repaired = repair_json(obj, return_objects=True)
        if isinstance(repaired, dict):
            return repaired
    except Exception:
        pass

    # Stage 5: 修复裸控制字符后再解析
    fixed = _fix_control_chars(obj)
    return json.loads(fixed)


async def _llm_parse(raw_text: str, model: str) -> dict:
    client = AsyncAnthropic(api_key=config.API_KEY, base_url=config.BASE_URL or None)
    user_prompt = (
        "请解析以下简历，返回指定 JSON 结构（未找到的字段返回 null 或 []）：\n\n"
        f"{raw_text}\n\n"
        f"JSON 格式示例：\n{_SCHEMA_EXAMPLE}"
    )

    last_err: Exception | None = None
    resp = None
    for attempt in range(3):
        try:
            resp = await client.messages.create(
                model=model,
                max_tokens=4096,
                temperature=0,
                system=_SYSTEM,
                messages=[{"role": "user", "content": user_prompt}],
            )
            if not resp.content:
                raise ValueError(f"API 返回空 content（stop_reason={resp.stop_reason}）")
            raw_resp = resp.content[0].text.strip()
            if not raw_resp:
                raise ValueError("API 返回空文本")
            parsed = _parse_llm_response(raw_resp)
            parsed["_model_used"] = model
            parsed["_raw_json"] = parsed.copy()
            return _normalize_salary(parsed)
        except Exception as e:
            last_err = e
            raw_start = "—"
            if resp and resp.content:
                raw_start = (resp.content[0].text or "")[:120]
            log.warning("LLM parse attempt %d failed: %s | raw_start=%.120s",
                        attempt + 1, e, raw_start)
            resp = None
            if attempt < 2:
                await asyncio.sleep(2 ** attempt)

    raise RuntimeError(f"LLM 解析失败（3次重试）: {last_err}") from last_err


# ── 主入口 ────────────────────────────────────────────────────────────────────

async def parse_and_save(
    resume_id: int,
    user_id: str,
    file_path: str,
    file_type: str,
) -> None:
    try:
        await resume_db.update_resume_status(resume_id, "parsing")

        raw_text = await _extract_text_async(file_path, file_type)
        await resume_db.update_resume_status(resume_id, "parsing", raw_text=raw_text)

        model = config.RESUME_PARSE_MODEL
        parsed = await _llm_parse(raw_text, model)

        await resume_db.save_parsed_result(resume_id, user_id, parsed)
        await resume_db.update_resume_status(resume_id, "done")
        log.info("Resume parsed: resume_id=%d user_id=%s", resume_id, user_id)

    except Exception as e:
        log.exception("Resume parse failed: resume_id=%d user_id=%s", resume_id, user_id)
        await resume_db.update_resume_status(resume_id, "failed", error=str(e))
